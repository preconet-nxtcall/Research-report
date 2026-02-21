import os
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
from backend.models import ResearchReport
from jinja2 import Template

DISCLAIMER = "This terminal uses AI-simulated data for illustrative research only. It does not use real-time market data and is not investment advice."
FOOTER_TEXT = "UNLISTED O’RENA – Where Opportunities Thrive in Private Markets | 10/2, Hungerford Street, Kolkata – 700017 | +91 90513 87093 | +91 98310 21208 | madhurr@unlistedorena.com | www.unlistedorena.com"

def _add_list_to_docx(doc: Document, items: list[str]):
    if not items:
        doc.add_paragraph("Data not provided.")
        return
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def generate_docx(report: ResearchReport, company_name: str) -> BytesIO:
    doc = Document()
    
    title = doc.add_heading(f"UNLISTED O’RENA TERMINAL REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(f"Secondary Share Investment Pitch: {company_name}", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Executive Summary
    doc.add_heading("1. EXECUTIVE SUMMARY", level=2)
    doc.add_paragraph(report.executiveSummary)

    # 2. Business Overview
    doc.add_heading("2. BUSINESS OVERVIEW & MARKET POSITIONING", level=2)
    doc.add_paragraph(report.businessOverviewAndMarketPositioning)

    # 3. Historical Financials
    doc.add_heading("3. HISTORICAL FINANCIAL PERFORMANCE (INR Cr)", level=2)
    t1 = doc.add_table(rows=1, cols=5)
    t1.style = 'Table Grid'
    h1 = t1.rows[0].cells
    for i, h in enumerate(["Year", "Revenue", "EBITDA", "PAT", "EPS"]):
        h1[i].text = h
        h1[i].paragraphs[0].runs[0].font.bold = True
    for item in report.historicalFinancialPerformance:
        row = t1.add_row().cells
        row[0].text = str(item.year)
        row[1].text = str(item.revenueINR_cr)
        row[2].text = str(item.ebitdaINR_cr)
        row[3].text = str(item.patINR_cr)
        row[4].text = str(item.epsINR)
        
    doc.add_paragraph()

    # 4. Financial Ratios
    doc.add_heading("4. FINANCIAL RATIOS & VALUATION METRICS", level=2)
    t2 = doc.add_table(rows=1, cols=6)
    t2.style = 'Table Grid'
    h2 = t2.rows[0].cells
    for i, h in enumerate(["P/E", "P/B", "EV/EBITDA", "EV/Sales", "ROE (%)", "ROCE (%)"]):
        h2[i].text = h
        h2[i].paragraphs[0].runs[0].font.bold = True
    row = t2.add_row().cells
    ratios = report.financialRatiosAndValuationMetrics
    row[0].text = str(ratios.peRatio)
    row[1].text = str(ratios.pbRatio)
    row[2].text = str(ratios.evToEbitda)
    row[3].text = str(ratios.evToSales)
    row[4].text = str(ratios.roePercent)
    row[5].text = str(ratios.rocePercent)
    
    doc.add_paragraph()

    # 5. Peer Comparison
    doc.add_heading("5. PEER COMPARISON ANALYSIS", level=2)
    t3 = doc.add_table(rows=1, cols=6)
    t3.style = 'Table Grid'
    h3 = t3.rows[0].cells
    for i, h in enumerate(["Company", "Revenue (INR Cr)", "EV/EBITDA", "P/E Ratio", "Rev Growth (%)", "ROE (%)"]):
        h3[i].text = h
        h3[i].paragraphs[0].runs[0].font.bold = True
    for p in report.peerComparisonAnalysis:
        row = t3.add_row().cells
        row[0].text = str(p.company)
        row[1].text = str(p.revenueINR_cr)
        row[2].text = str(p.evToEbitda)
        row[3].text = str(p.peRatio)
        row[4].text = str(p.revenueGrowthPercent)
        row[5].text = str(p.roePercent)

    doc.add_paragraph()

    # 6. Unit Economics
    doc.add_heading("6. UNIT ECONOMICS & KPIs", level=2)
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = 'Table Grid'
    h4 = t4.rows[0].cells
    h4[0].text, h4[1].text = "Metric", "Value"
    h4[0].paragraphs[0].runs[0].font.bold = True
    h4[1].paragraphs[0].runs[0].font.bold = True
    for item in report.unitEconomicsAndKPIs:
        row = t4.add_row().cells
        row[0].text = str(item.metric)
        row[1].text = str(item.value)

    doc.add_paragraph()

    # 7. Capital Structure
    doc.add_heading("7. CAPITAL STRUCTURE & SECONDARY OFFER", level=2)
    cap = report.capitalStructureAndSecondaryOffer
    t5 = doc.add_table(rows=1, cols=2)
    t5.style = 'Table Grid'
    h5 = t5.rows[0].cells
    h5[0].text, h5[1].text = "Detail", "Value"
    h5[0].paragraphs[0].runs[0].font.bold = True
    h5[1].paragraphs[0].runs[0].font.bold = True
    mapping = [
        ("Total Shares Outstanding (mn)", str(cap.totalSharesOutstanding_mn)),
        ("Promoter Holding (%)", str(cap.promoterHoldingPercent)),
        ("Institutional Holding (%)", str(cap.institutionalHoldingPercent)),
        ("Employee Holding (%)", str(cap.employeeHoldingPercent)),
        ("Indicative Valuation (INR Cr)", str(cap.indicativeValuationINR_cr)),
        ("Secondary Offer Size (INR Cr)", str(cap.secondaryOfferSizeINR_cr)),
        ("Implied Per Share Value (INR)", str(cap.impliedPerShareValueINR)),
    ]
    for m in mapping:
        row = t5.add_row().cells
        row[0].text = m[0]
        row[1].text = m[1]

    doc.add_paragraph()

    # 8. Projections
    doc.add_heading("8. FORWARD PROJECTIONS", level=2)
    t6 = doc.add_table(rows=1, cols=4)
    t6.style = 'Table Grid'
    h6 = t6.rows[0].cells
    for i, h in enumerate(["Scenario", "Revenue (INR Cr)", "EBITDA Margin (%)", "PAT (INR Cr)"]):
        h6[i].text = h
        h6[i].paragraphs[0].runs[0].font.bold = True
    for proj in report.forwardProjections:
        row = t6.add_row().cells
        row[0].text = str(proj.scenario)
        row[1].text = str(proj.revenueINR_cr)
        row[2].text = str(proj.ebitdaMarginPercent)
        row[3].text = str(proj.patINR_cr)

    doc.add_paragraph()
    
    # 9. Exit Pathways
    doc.add_heading("9. EXIT PATHWAYS & LIQUIDITY", level=2)
    doc.add_paragraph(report.exitPathwaysAndLiquidity)
    
    # 10. SWOT
    doc.add_heading("10. SWOT ANALYSIS", level=2)
    doc.add_heading("Strengths", level=3)
    _add_list_to_docx(doc, report.swotAnalysis.strengths)
    doc.add_heading("Weaknesses", level=3)
    _add_list_to_docx(doc, report.swotAnalysis.weaknesses)
    doc.add_heading("Opportunities", level=3)
    _add_list_to_docx(doc, report.swotAnalysis.opportunities)
    doc.add_heading("Threats", level=3)
    _add_list_to_docx(doc, report.swotAnalysis.threats)
    
    # 11. Governance
    doc.add_heading("11. GOVERNANCE, COMPLIANCE & RISK", level=2)
    _add_list_to_docx(doc, report.governanceComplianceAndRisk)

    # 12. Highlights
    doc.add_heading("12. INVESTMENT HIGHLIGHTS & THESIS", level=2)
    _add_list_to_docx(doc, report.investmentHighlightsAndThesis)

    # 13. Scorecard
    doc.add_heading("13. INVESTMENT SCORECARD", level=2)
    t7 = doc.add_table(rows=1, cols=2)
    t7.style = 'Table Grid'
    h7 = t7.rows[0].cells
    h7[0].text, h7[1].text = "Parameter", "Rating (/5)"
    h7[0].paragraphs[0].runs[0].font.bold = True
    h7[1].paragraphs[0].runs[0].font.bold = True
    for item in report.investmentScorecard:
        row = t7.add_row().cells
        row[0].text = str(item.parameter)
        row[1].text = str(item.ratingOutOf5)
        
    doc.add_paragraph()
    
    # 14. sizing
    doc.add_heading("14. INVESTMENT SIZING & MINIMUMS", level=2)
    doc.add_paragraph(report.investmentSizingAndMinimums)

    # 15. IRR/MOIC
    doc.add_heading("15. VALUATION SENSITIVITY & ASSUMPTIONS", level=2)
    sens = report.assumptionsAndSensitivityAnalysis
    t8 = doc.add_table(rows=1, cols=2)
    t8.style = 'Table Grid'
    h8 = t8.rows[0].cells
    h8[0].text, h8[1].text = "Parameter", "Assumption Value"
    h8[0].paragraphs[0].runs[0].font.bold = True
    h8[1].paragraphs[0].runs[0].font.bold = True
    smapping = [
        ("Entry Valuation (INR Cr)", str(sens.entryValuationINR_cr)),
        ("Exit Valuation (INR Cr)", str(sens.exitValuationINR_cr)),
        ("Holding Period (Years)", str(sens.holdingPeriod_years)),
        ("Base Case IRR (%)", str(sens.baseCaseIRR_percent)),
        ("Bull Case IRR (%)", str(sens.bullCaseIRR_percent)),
        ("Bear Case IRR (%)", str(sens.bearCaseIRR_percent)),
        ("MOIC (Base)", str(sens.moic_base)),
        ("Exit Multiple Assumption", sens.exitMultipleAssumption),
    ]
    for m in smapping:
        row = t8.add_row().cells
        row[0].text = m[0]
        row[1].text = m[1]
    
    doc.add_paragraph(f"Commentary: {sens.valuationSensitivityCommentary}")
    doc.add_paragraph()
    
    # 16. Dataroom
    doc.add_heading("16. DATA ROOM & DUE DILIGENCE", level=2)
    _add_list_to_docx(doc, report.dataRoomAndDueDiligenceChecklist)
    
    # 17. Timeline
    doc.add_heading("17. NEXT STEPS & TIMELINE", level=2)
    doc.add_paragraph(report.nextStepsAndTransactionTimeline)

    doc.add_heading("18. ASSUMPTIONS & LIMITATIONS", level=2)
    doc.add_paragraph(report.assumptionsAndLimitations)

    # Footer and Disclaimer
    doc.add_paragraph("-" * 80)
    footer_p1 = doc.add_paragraph()
    footer_p1.add_run(DISCLAIMER).italic = True
    
    footer_p2 = doc.add_paragraph()
    footer_p2.add_run(FOOTER_TEXT).bold = True
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(report: ResearchReport, company_name: str) -> BytesIO:
    html_template = """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        @page {
            size: a4;
            margin: 2cm;
            @frame footer {
                -pdf-frame-content: footer_content;
                bottom: 1cm;
                margin-left: 2cm;
                margin-right: 2cm;
                height: 2cm;
            }
        }
        body { font-family: Helvetica, Arial, sans-serif; font-size: 11px; line-height: 1.5; color: #1e293b; }
        h1 { font-size: 24px; color: #020617; text-align: center; border-bottom: 2px solid #000; padding-bottom: 10px; margin-bottom: 20px;}
        h2 { font-size: 16px; color: #0f172a; margin-top: 20px; border-bottom: 1px solid #cbd5e1; padding-bottom: 5px;}
        h3 { font-size: 14px; color: #334155; margin-top: 15px;}
        p { margin-bottom: 10px; }
        ul { margin-bottom: 10px; }
        li { margin-bottom: 5px; }
        table { width: 100%; border-collapse: collapse; margin-bottom: 15px; font-size: 10px; }
        th, td { border: 1px solid #cbd5e1; padding: 6px; text-align: left; }
        th { background-color: #f1f5f9; font-weight: bold; }
        .disclaimer { font-size: 10px; color: #64748b; font-style: italic; margin-bottom: 5px; text-align: center; border-top: 1px solid #ccc; padding-top: 10px;}
        .footer-text { font-size: 9px; color: #020617; font-weight: bold; text-align: center; }
    </style>
    </head>
    <body>
        <h1>UNLISTED O’RENA TERMINAL</h1>
        <p style="text-align: center; font-size: 16px; font-weight: bold;">Secondary Share Investment Pitch: {{ company_name }}</p>
        
        <h2>1. EXECUTIVE SUMMARY</h2>
        <p>{{ report.executiveSummary | replace('\\n', '<br>') }}</p>

        <h2>2. BUSINESS OVERVIEW & MARKET POSITIONING</h2>
        <p>{{ report.businessOverviewAndMarketPositioning | replace('\\n', '<br>') }}</p>

        <h2>3. HISTORICAL FINANCIAL PERFORMANCE (INR Cr)</h2>
        <table>
            <thead><tr><th>Year</th><th>Revenue</th><th>EBITDA</th><th>PAT</th><th>EPS</th></tr></thead>
            <tbody>
            {% for row in report.historicalFinancialPerformance %}
                <tr><td>{{ row.year }}</td><td>{{ row.revenueINR_cr }}</td><td>{{ row.ebitdaINR_cr }}</td><td>{{ row.patINR_cr }}</td><td>{{ row.epsINR }}</td></tr>
            {% endfor %}
            </tbody>
        </table>

        <h2>4. FINANCIAL RATIOS & VALUATION METRICS</h2>
        <table>
            <thead><tr><th>P/E</th><th>P/B</th><th>EV/EBITDA</th><th>EV/Sales</th><th>ROE (%)</th><th>ROCE (%)</th></tr></thead>
            <tbody>
                <tr>
                    <td>{{ report.financialRatiosAndValuationMetrics.peRatio }}</td>
                    <td>{{ report.financialRatiosAndValuationMetrics.pbRatio }}</td>
                    <td>{{ report.financialRatiosAndValuationMetrics.evToEbitda }}</td>
                    <td>{{ report.financialRatiosAndValuationMetrics.evToSales }}</td>
                    <td>{{ report.financialRatiosAndValuationMetrics.roePercent }}</td>
                    <td>{{ report.financialRatiosAndValuationMetrics.rocePercent }}</td>
                </tr>
            </tbody>
        </table>

        <h2>5. PEER COMPARISON ANALYSIS</h2>
        <table>
            <thead><tr><th>Company</th><th>Revenue (INR Cr)</th><th>EV/EBITDA</th><th>P/E Ratio</th><th>Rev Growth (%)</th><th>ROE (%)</th></tr></thead>
            <tbody>
            {% for peer in report.peerComparisonAnalysis %}
                <tr><td>{{ peer.company }}</td><td>{{ peer.revenueINR_cr }}</td><td>{{ peer.evToEbitda }}</td><td>{{ peer.peRatio }}</td><td>{{ peer.revenueGrowthPercent }}</td><td>{{ peer.roePercent }}</td></tr>
            {% endfor %}
            </tbody>
        </table>

        <h2>6. UNIT ECONOMICS & KPIs</h2>
        <table>
            <thead><tr><th>Metric</th><th>Value</th></tr></thead>
            <tbody>
            {% for item in report.unitEconomicsAndKPIs %}
                <tr><td>{{ item.metric }}</td><td>{{ item.value }}</td></tr>
            {% endfor %}
            </tbody>
        </table>

        <h2>7. CAPITAL STRUCTURE & SECONDARY OFFER</h2>
        <table>
            <thead><tr><th>Detail</th><th>Value</th></tr></thead>
            <tbody>
                <tr><td>Total Shares Outstanding (mn)</td><td>{{ report.capitalStructureAndSecondaryOffer.totalSharesOutstanding_mn }}</td></tr>
                <tr><td>Promoter Holding (%)</td><td>{{ report.capitalStructureAndSecondaryOffer.promoterHoldingPercent }}</td></tr>
                <tr><td>Institutional Holding (%)</td><td>{{ report.capitalStructureAndSecondaryOffer.institutionalHoldingPercent }}</td></tr>
                <tr><td>Employee Holding (%)</td><td>{{ report.capitalStructureAndSecondaryOffer.employeeHoldingPercent }}</td></tr>
                <tr><td>Indicative Valuation (INR Cr)</td><td>{{ report.capitalStructureAndSecondaryOffer.indicativeValuationINR_cr }}</td></tr>
                <tr><td>Secondary Offer Size (INR Cr)</td><td>{{ report.capitalStructureAndSecondaryOffer.secondaryOfferSizeINR_cr }}</td></tr>
                <tr><td>Implied Per Share Value (INR)</td><td>{{ report.capitalStructureAndSecondaryOffer.impliedPerShareValueINR }}</td></tr>
            </tbody>
        </table>

        <h2>8. FORWARD PROJECTIONS</h2>
        <table>
            <thead><tr><th>Scenario</th><th>Revenue (INR Cr)</th><th>EBITDA Margin (%)</th><th>PAT (INR Cr)</th></tr></thead>
            <tbody>
            {% for proj in report.forwardProjections %}
                <tr><td>{{ proj.scenario }}</td><td>{{ proj.revenueINR_cr }}</td><td>{{ proj.ebitdaMarginPercent }}</td><td>{{ proj.patINR_cr }}</td></tr>
            {% endfor %}
            </tbody>
        </table>

        <h2>9. EXIT PATHWAYS & LIQUIDITY</h2>
        <p>{{ report.exitPathwaysAndLiquidity | replace('\\n', '<br>') }}</p>

        <h2>10. SWOT ANALYSIS</h2>
        <h3>Strengths</h3>
        <ul>
        {% for s in report.swotAnalysis.strengths %}
            <li>{{ s }}</li>
        {% endfor %}
        </ul>
        <h3>Weaknesses</h3>
        <ul>
        {% for w in report.swotAnalysis.weaknesses %}
            <li>{{ w }}</li>
        {% endfor %}
        </ul>
        <h3>Opportunities</h3>
        <ul>
        {% for o in report.swotAnalysis.opportunities %}
            <li>{{ o }}</li>
        {% endfor %}
        </ul>
        <h3>Threats</h3>
        <ul>
        {% for t in report.swotAnalysis.threats %}
            <li>{{ t }}</li>
        {% endfor %}
        </ul>

        <h2>11. GOVERNANCE, COMPLIANCE & RISK</h2>
        <ul>
        {% for g in report.governanceComplianceAndRisk %}
            <li>{{ g }}</li>
        {% endfor %}
        </ul>

        <h2>12. INVESTMENT HIGHLIGHTS & THESIS</h2>
        <ul>
        {% for info in report.investmentHighlightsAndThesis %}
            <li>{{ info }}</li>
        {% endfor %}
        </ul>

        <h2>13. INVESTMENT SCORECARD</h2>
        <table>
            <thead><tr><th>Parameter</th><th>Rating (/5)</th></tr></thead>
            <tbody>
            {% for score in report.investmentScorecard %}
                <tr><td>{{ score.parameter }}</td><td>{{ score.ratingOutOf5 }}</td></tr>
            {% endfor %}
            </tbody>
        </table>

        <h2>14. INVESTMENT SIZING & MINIMUMS</h2>
        <p>{{ report.investmentSizingAndMinimums | replace('\\n', '<br>') }}</p>

        <h2>15. VALUATION SENSITIVITY & ASSUMPTIONS</h2>
        <table>
            <thead><tr><th>Parameter</th><th>Assumption Value</th></tr></thead>
            <tbody>
                <tr><td>Entry Valuation (INR Cr)</td><td>{{ report.assumptionsAndSensitivityAnalysis.entryValuationINR_cr }}</td></tr>
                <tr><td>Exit Valuation (INR Cr)</td><td>{{ report.assumptionsAndSensitivityAnalysis.exitValuationINR_cr }}</td></tr>
                <tr><td>Holding Period (Years)</td><td>{{ report.assumptionsAndSensitivityAnalysis.holdingPeriod_years }}</td></tr>
                <tr><td>Base Case IRR (%)</td><td>{{ report.assumptionsAndSensitivityAnalysis.baseCaseIRR_percent }}</td></tr>
                <tr><td>Bull Case IRR (%)</td><td>{{ report.assumptionsAndSensitivityAnalysis.bullCaseIRR_percent }}</td></tr>
                <tr><td>Bear Case IRR (%)</td><td>{{ report.assumptionsAndSensitivityAnalysis.bearCaseIRR_percent }}</td></tr>
                <tr><td>MOIC (Base)</td><td>{{ report.assumptionsAndSensitivityAnalysis.moic_base }}</td></tr>
                <tr><td>Exit Multiple Assumption</td><td>{{ report.assumptionsAndSensitivityAnalysis.exitMultipleAssumption }}</td></tr>
            </tbody>
        </table>
        <p><strong>Valuation Commentary:</strong><br>{{ report.assumptionsAndSensitivityAnalysis.valuationSensitivityCommentary | replace('\\n', '<br>') }}</p>

        <h2>16. DATA ROOM & DUE DILIGENCE</h2>
        <ul>
        {% for d in report.dataRoomAndDueDiligenceChecklist %}
            <li>{{ d }}</li>
        {% endfor %}
        </ul>

        <h2>17. NEXT STEPS & TIMELINE</h2>
        <p>{{ report.nextStepsAndTransactionTimeline | replace('\\n', '<br>') }}</p>

        <h2>18. ASSUMPTIONS & LIMITATIONS</h2>
        <p>{{ report.assumptionsAndLimitations | replace('\\n', '<br>') }}</p>

        <div id="footer_content">
            <div class="disclaimer">{{ disclaimer }}</div>
            <div class="footer-text">{{ footer_text }}</div>
        </div>
    </body>
    </html>
    """

    template = Template(html_template)
    html_content = template.render(
        report=report, 
        company_name=company_name,
        disclaimer=DISCLAIMER,
        footer_text=FOOTER_TEXT
    )

    buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=buffer)
    
    if pisa_status.err:
        raise Exception("Error generating PDF")
        
    buffer.seek(0)
    return buffer
